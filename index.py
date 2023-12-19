import os
import requests
from bs4 import BeautifulSoup
from docx import Document

def download_xml_content(xml_url):
    response = requests.get(xml_url)
    return response.text

def extract_and_filter_urls(xml_content, exclude_keywords):
    soup = BeautifulSoup(xml_content, 'xml')
    urls = [element.text for element in soup.find_all('loc')]
    filtered_urls = [url for url in urls if all(keyword not in url for keyword in exclude_keywords)]
    return filtered_urls

def main(xml_url, exclude_keywords):
    xml_content = download_xml_content(xml_url)
    filtered_urls = extract_and_filter_urls(xml_content, exclude_keywords)
    return filtered_urls

# Example usage
xml_url = 'https://allabouttattoo.com/post-sitemap.xml'  # Replace with the actual XML URL
exclude_keywords = ['.jpg', '.png', '.js', 'wp-content']
urls = main(xml_url, exclude_keywords)



def fetch_page_content(url):
    try:
        response = requests.get(url)
        return response.content
    except requests.RequestException as e:
        print(f"Error fetching {url}: {e}")
        return None

def save_content_to_doc(content, filename, folder="downloaded_docs"):
    # Ensure folder exists
    if not os.path.exists(folder):
        os.makedirs(folder)

    doc = Document()
    doc.add_heading(filename, level=1)
    doc.add_paragraph(content)
    filepath = os.path.join(folder, f"{filename}.docx")
    doc.save(filepath)

def main(urls):
    for url in urls:
        print(f"Processing {url}")
        content = fetch_page_content(url)
        if content:
            soup = BeautifulSoup(content, 'html.parser')
            text_content = soup.get_text()  # Extracts text; customize as needed
            save_content_to_doc(text_content, url.split('/')[3])  # Filename from URL

# Example usage
# urls=["https://allabouttattoo.com/neck-lip-tattoo/","https://allabouttattoo.com/do-finger-tattoos-hurt/"]
main(urls)
